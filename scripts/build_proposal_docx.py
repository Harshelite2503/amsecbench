"""Build docs/AMSecBench_Research_Proposal.docx from live benchmark results."""
import json
import pandas as pd
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from amsec.config import RENDER_DIR, RESULTS_DIR, ROOT

man = pd.read_csv(RESULTS_DIR / "manifest.csv"); rules = pd.read_csv(RESULTS_DIR / "rules_detector.csv")
ml = pd.read_csv(RESULTS_DIR / "ml_detectors.csv"); dm = json.load(open(RESULTS_DIR / "detector_matrix.json"))
atk = dm["attacks"]; n_files = len(man); n_clean = int((man.attack == "none").sum()); n_att = n_files - n_clean
rules_fpr = rules[rules.attack == "none"].flagged.mean(); rules_by = {a: rules[rules.attack == a].flagged.mean() for a in atk}
best = ml.sort_values("auc", ascending=False).iloc[0]

doc = Document(); doc.styles["Normal"].font.name = "Calibri"; doc.styles["Normal"].font.size = Pt(11)
for s in doc.sections: s.left_margin = s.right_margin = s.top_margin = s.bottom_margin = Inches(1)
def H(t, l=1): doc.add_heading(t, level=l)
def P(t, bold=False, italic=False, align=None):
    p = doc.add_paragraph(); r = p.add_run(t); r.bold = bold; r.italic = italic
    if align is not None: p.alignment = align
def B(items, style="List Bullet"):
    for i in items: doc.add_paragraph(i, style=style)
def T(header, body, widths=None):
    t = doc.add_table(rows=1, cols=len(header)); t.style = "Light Grid Accent 1"; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(header):
        c = t.rows[0].cells[i]; c.text = ""; c.paragraphs[0].add_run(h).bold = True
    for row in body:
        cells = t.add_row().cells
        for i, v in enumerate(row): cells[i].text = str(v)
    if widths:
        for row in t.rows:
            for i, w in enumerate(widths): row.cells[i].width = Inches(w)
    doc.add_paragraph()
def FIG(path, cap, w=6.0):
    doc.add_picture(str(path), width=Inches(w)); doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    P(cap, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER)

p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("AMSecBench"); r.bold = True; r.font.size = Pt(24)
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("AI-Native Security for Additive Manufacturing: Can Large Language and Vision-Language Models Detect Toolpath Sabotage — and Can They Defeat Design Obfuscation?"); r.font.size = Pt(14); r.italic = True
P("Research Proposal", align=WD_ALIGN_PARAGRAPH.CENTER)
P("Harsh Vardhan Gupta (proposer)  ·  Prof. Nikhil Gupta, FASM  ·  Prof. Ramesh Karri\nNYU Tandon School of Engineering / NYU Center for Cybersecurity", align=WD_ALIGN_PARAGRAPH.CENTER)
P("August 2026  ·  Code: github.com/Harshelite2503/amsecbench", italic=True, align=WD_ALIGN_PARAGRAPH.CENTER)

H("Abstract")
P(f"Additive manufacturing (AM) is a cyber-physical process: a design file becomes a toolpath (G-code) becomes a part, and every hop is a place where an attacker can degrade the part without changing how it looks. The AM-security literature — including the Proc. IEEE survey by Gupta, Karri and colleagues — established the attack taxonomy (internal voids, infill reduction, parameter manipulation, dimensional sabotage) and two defensive families: verification of the toolpath against the design, and design obfuscation so a stolen file prints the wrong part. Large language and vision-language models (LLMs/VLMs) have not been evaluated on either side of this problem. We propose AMSecBench, a fully synthetic, labelled benchmark with two tracks. Track A (defense) asks whether LLMs/VLMs can detect G-code sabotage, with and without a trusted reference, better than rule-based and classical-ML detectors. Track B (red team) asks whether LLMs/VLMs can defeat the design-based obfuscation schemes developed in the collaborating group by recovering the intended design from an obfuscated file. A working benchmark exists: {n_files} G-code files ({n_clean} clean incl. benign parameter variants, {n_att} attacked across 7 attack classes on 5 parts). Baselines already show the central difficulty: a reference-comparison detector catches 75–100% of attacks but falsely flags {rules_fpr:.0%} of legitimate parameter variants, and reference-free ML reaches AUC {best.auc:.2f} with a {best.false_positive_rate:.0%} false-positive rate. Outcomes are a benchmark paper at a security venue, a defense/attack findings paper, and a proposal-ready basis for NSF SaTC / DoD funding.")

H("1. Problem Statement")
H("1.1 The AM pipeline is attackable at every hop and the damage is invisible", 2)
P("A 3D-printed part is produced from a CAD/STL file through slicer software into G-code executed by the printer. Prior work (dr0wned, 2017; Gupta et al., 2020; Rossel et al., USENIX Security 2025) has shown that small edits to the G-code — removing a few millimetres of extrusion to create a void, reducing infill in a band of layers, lowering nozzle temperature for a few layers — yield parts that pass visual inspection and fail in service. Distributed and outsourced printing widens the attack surface: the file leaves the designer’s control.")
H("1.2 Existing detectors are brittle", 2)
P("Verification approaches compare the toolpath to the design or to a trusted reference. They work when the reference and the suspect were produced with identical settings, but legitimate operators change infill, layer height, temperature and speed routinely; a naive comparison flags those changes as attacks. Reference-free anomaly detectors must instead generalise across geometries they have never seen. Neither approach reasons about intent — whether a change is a plausible operator choice or a degradation.")
H("1.3 Obfuscation defenses have never faced an AI adversary", 2)
P("Design-based security (Chen & Gupta; NYU) embeds decoy features and parameter-gated ‘key’ features so a stolen STL prints a wrong or failed part unless the process recipe is known. These schemes were evaluated against human reverse engineers and conventional tooling. Multimodal LLMs can now read STL geometry, renders and G-code, and reason about mechanical function; whether they can identify and strip decoys is unknown — and directly determines whether these defenses remain valid.")
H("1.4 LLMs are entering manufacturing workflows without a security evaluation", 2)
P("LLM assistants are being integrated into slicers, print-farm management and CAD tools. There is no benchmark that measures their value as a detector or their threat as an attacker in AM. Both the defensive community and the roadmap literature (Advanced Intelligent Systems 2026) identify this as an open problem.")

H("2. Objectives")
B(["O1 — Benchmark: an open, synthetic, reproducible dataset of clean, benign-variant and sabotaged G-code with ground-truth labels and attack metadata, plus rendered toolpath images.",
   "O2 — Defense evaluation: measure LLM (text) and VLM (image) detection accuracy, false-positive rate on benign variants, attack-class attribution and calibration, with and without a trusted reference, against rule and ML baselines.",
   "O3 — Red-team evaluation: implement the group’s obfuscation schemes on benchmark parts and measure whether LLMs/VLMs recover the intended design (IoU) and identify parameter-gated features.",
   "O4 — Hardening: use red-team results to propose obfuscation variants that resist multimodal AI, and LLM-in-the-loop detection that reasons about operator intent.",
   "O5 — Real-slicer validation: reproduce the key results on PrusaSlicer/Cura output and, where available, on printed specimens from the collaborating lab."], "List Number")

H("3. Background and Related Work")
P("AM security surveys: Gupta, Karri et al. (Proc. IEEE 2020) taxonomise threats and countermeasures including obfuscation and embedded security features. Attack studies: dr0wned (Belikovetsky et al. 2017) demonstrated void-sabotage of a drone propeller; Rossel et al. (USENIX Security 2025) analysed malicious G-code at scale across slicers. Detection: probabilistic anomaly models on printer telemetry (Yoginath et al. 2022), power side-channels (PowerGuard 2024), design-vs-toolpath verification (Mahmood et al. 2025), hardware firewalls (Firewall3D 2026). Design-based security: Chen (PhD, NYU 2019, advised by N. Gupta) — embedded tracking codes and parameter-gated features. AI in AM: roadmap papers call for LLM/VLM quality assurance but no security benchmark exists.")

H("4. Approach")
H("4.1 Track A — sabotage detection benchmark", 2)
T(["Component", "Design", "Status"], [
 ["Parts", "5 parametric parts: block, bracket with holes, thin-wall tube, ASTM-style dog-bone, grid plate (generated, no downloads)", "Done"],
 ["Slicer", "Minimal inspectable FDM slicer (perimeters + rectilinear infill, absolute E) so every attack has a known effect", "Done"],
 ["Attacks", "infill reduction, void insertion (segment-split), layer-height change, mid-print temperature drop, 3% XY scaling, under-extrusion band, infill toolpath jitter", "Done"],
 ["Benign variants", "8 legitimate parameter changes per part (infill 0.2/0.4, layer 0.15/0.25, temp 200/220, speed 30/60) — detectors must not flag these", "Done"],
 ["Features / renders", "17 per-file features (layer statistics, extrusion-rate, infill ratio, temperature, bbox, jitter, travel-in-fill); PNG toolpath renders", "Done"],
 ["Baselines", "Rule detector (reference comparison); logistic / random-forest / gradient-boosting, leave-one-part-out", "Done"],
 ["LLM / VLM detectors", "Claude on per-layer summaries with reference; on raw G-code without reference; on rendered images", "Code ready; runs pending"],
], widths=[1.3, 3.9, 1.3])
H("4.2 Track B — obfuscation red team", 2)
B(["Decoy features: extra bosses added to the intended design; the attacker must decide which features to remove.",
   "Parameter-gated keys: 0.25 mm ribs that only print with a specific nozzle profile and carry a load path.",
   "Adversary: LLM/VLM given STL statistics, multi-view renders and (optionally) the G-code, asked to output the intended design as a feature list / edit script; scored by volumetric IoU against ground truth and by whether keys are identified.",
   "Threat model: attacker has the file but not the process recipe; no physical access; this mirrors the design-theft scenario the defenses target."])
H("4.3 Ethics and scope", 2)
P("All artefacts are synthetic and generated locally. No firmware, network or physical printer is targeted, and no exploit tooling against real systems is produced. Track B evaluates academic defenses published by the collaborating group, with their authors, which is standard practice for security research.")

H("5. Preliminary Results")
P(f"The Track A benchmark is built and two baseline families have been evaluated on {n_files} G-code files ({n_clean} clean/benign, {n_att} attacked; attacks that had no geometric effect, e.g. a void placed in a part with no infill, are excluded rather than mislabelled).")
FIG(RESULTS_DIR / "detector_matrix.png", "Figure 1. Detection rate per attack class for the baseline detectors; the last column is the false-positive rate on benign parameter variants.")
H("5.1 Reference comparison is sensitive but not specific", 2)
P("The rule detector, which compares a suspect file with the trusted reference, detects " + ", ".join(f"{a.replace('_', ' ')} {v:.0%}" for a, v in rules_by.items()) + f". However it flags {rules_fpr:.0%} of legitimate parameter variants as attacks — every layer-height and temperature change an operator might legitimately make is indistinguishable from sabotage without reasoning about intent. This is the specificity problem that motivates an LLM-in-the-loop detector.")
H("5.2 Reference-free ML generalises only partially to unseen parts", 2)
T(["Model", "AUC", "F1", "False-positive rate"], [[m.model, m.auc, m.f1, m.false_positive_rate] for _, m in ml.iterrows()], widths=[2.0, 1.0, 1.0, 1.8])
P(f"Under leave-one-part-out evaluation (the model never sees the test geometry), the best model reaches AUC {best.auc:.2f} but still mislabels {best.false_positive_rate:.0%} of benign files. Void insertion and infill reduction are the hardest classes because their signature is local (a few layers) while the features are global. These are exactly the cases where a model that can read the per-layer table and reason about where the change occurred should help.")
H("5.3 Engineering findings", 2)
B(["A void attack must split toolpath segments at the void boundary; naïve endpoint deletion is a no-op on edge-to-edge infill lines — a subtlety that likely affects earlier synthetic studies.",
   "Per-layer features are required to see small voids; whole-file ratios dilute them below the benign-variation noise floor (benign variants move the travel-in-fill feature by ≤0.024, real voids by 0.05–0.32).",
   "Thin-wall parts carry no infill, so void/infill attacks are inapplicable — the benchmark records this rather than mislabelling."])
FIG(RENDER_DIR / "bracket_void_insertion_s0.png", "Figure 2. Rendered toolpath of the bracket with an inserted void (right: the affected layer).", 5.2)

H("6. Work Plan")
T(["Phase", "Weeks", "Activities", "Deliverable"], [
 ["1. LLM/VLM detection", "1–3", "Run Claude detectors (summary+reference, raw no-reference, image); calibration; attribution accuracy; cost", "Track A results"],
 ["2. Real-slicer port", "3–5", "Regenerate attacks on PrusaSlicer/Cura G-code; confirm findings transfer", "Benchmark v1.0"],
 ["3. Obfuscation red team", "4–8", "Implement Chen–Gupta schemes on benchmark parts; LLM/VLM de-obfuscation; IoU scoring; hardening variants", "Track B results"],
 ["4. Physical validation", "6–10", "Print a subset (clean vs sabotaged) in the Gupta lab; µCT/mechanical test to confirm attack effect", "Validation data"],
 ["5. Writing / proposal", "8–12", "Benchmark + findings papers; NSF SaTC / DoD white paper", "Submissions"],
], widths=[1.4, 0.7, 3.2, 1.2])

H("7. Expected Outputs")
T(["Output", "Venue options", "Contribution"], [
 ["Benchmark + detection paper", "USENIX Security / IEEE S&P (dataset & benchmark track); ACSAC; Additive Manufacturing", "First LLM/VLM evaluation for AM sabotage detection with benign-variant false-positive analysis"],
 ["Red-team paper", "IEEE S&P / CCS workshops (CPS-SPC); Journal of Manufacturing Systems", "Whether design obfuscation survives multimodal AI; hardened variants"],
 ["Open benchmark + code", "GitHub + Zenodo", "Reusable, extendable to other AM processes"],
 ["Funding proposal", "NSF SaTC, DoD/ARL manufacturing security, NYU CCS seed", "Both PIs have track record; benchmark provides preliminary data"],
], widths=[1.7, 2.3, 2.5])

H("8. Resources and Budget")
T(["Item", "Estimate"], [
 ["LLM/VLM detection runs (Track A, ~150 files × 3 modes)", "≈ US$50–100"],
 ["Red-team runs (Track B, 10 parts × 2 schemes × repeats)", "≈ US$30–60"],
 ["Printing + testing of validation specimens", "Gupta lab consumables; µCT time"],
 ["Compute", "Laptop-scale"],
], widths=[4.5, 2.0])

H("9. Roles")
B(["Harsh Vardhan Gupta: benchmark engineering, attack/detector implementation, LLM/VLM experiments, analysis, drafting.",
   "Prof. Nikhil Gupta: AM process and materials expertise, obfuscation schemes, physical validation, co-authorship.",
   "Prof. Ramesh Karri: threat modelling, security-venue framing, red-team methodology, co-authorship and funding strategy."])

H("10. Risks and Mitigations")
T(["Risk", "Mitigation"], [
 ["Synthetic slicer differs from real slicers", "Phase 2 port to PrusaSlicer/Cura; attacks operate on G-code so they transfer"],
 ["LLM cost / context limits on raw G-code", "Per-layer summary mode; chunked windows; Message Batches"],
 ["Dual-use concern in Track B", "Synthetic parts only; academic defenses evaluated with their authors; no real-system exploitation; responsible disclosure norms"],
 ["Small part library", "Parametric generator scales to hundreds of parts; add open-source functional parts"],
], widths=[2.6, 3.9])

H("References")
B(["Gupta N., Karri R. et al. A survey of cybersecurity of digital manufacturing. Proceedings of the IEEE, 2020.",
   "Chen F. Design-based security strategies for the additive manufacturing cyber-physical systems. PhD thesis, NYU, 2019.",
   "Belikovetsky S. et al. dr0wned — cyber-physical attack with additive manufacturing. USENIX WOOT, 2017.",
   "Rossel J. et al. Security implications of malicious G-codes in 3D printing. USENIX Security, 2025.",
   "Yoginath S. et al. Stealthy cyber anomaly detection on large noisy multi-material 3D printer datasets. ACM AMSec, 2022.",
   "Mahmood M.A. et al. A novel framework for identification of cyber-physical attacks in additive manufacturing. Progress in Additive Manufacturing, 2025.",
   "Asgar S.A.G., Reddy N. Firewall3D: a hardware firewall for defending 3D printers against firmware attacks. arXiv, 2026.",
   "Zolfagharian A. et al. Roadmap on artificial intelligence-augmented additive manufacturing. Advanced Intelligent Systems, 2026.",
   "Gupta N., Beckwith C. Architected Metamaterials: Design Principles and Properties. Springer, 2025."])

out = ROOT / "AMSecBench_Research_Proposal.docx"; doc.save(out); (ROOT / "docs" / out.name).write_bytes(out.read_bytes()); print("saved", out)
